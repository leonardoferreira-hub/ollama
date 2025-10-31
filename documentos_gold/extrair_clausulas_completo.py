#!/usr/bin/env python3
"""
Script completo para extrair e catalogar cláusulas dos documentos GOLD
"""
import json
from docx import Document
from pathlib import Path
import re

def extrair_estrutura_completa(caminho_docx):
    """Extrai a estrutura completa do documento incluindo estilos"""
    doc = Document(caminho_docx)
    estrutura = []
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            estrutura.append({
                'indice': i,
                'texto': para.text.strip(),
                'estilo': para.style.name if para.style else 'Normal',
                'nivel': para.style.name.replace('Heading ', '') if 'Heading' in para.style.name else None
            })
    
    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    estrutura.append({
                        'indice': len(estrutura),
                        'texto': cell.text.strip(),
                        'estilo': 'Tabela',
                        'nivel': None
                    })
    
    return estrutura

def processar_clausulas_detalhadas(estrutura):
    """Processa e identifica todas as cláusulas com mais precisão"""
    clausulas = []
    
    # Padrões mais abrangentes
    padroes = [
        re.compile(r'^(\d+)\.\s+([A-Z][^.]+)'),  # 1. TÍTULO
        re.compile(r'^(\d+\.\d+)\.\s+(.+)'),      # 1.1. Subtítulo
        re.compile(r'^(\d+\.\d+\.\d+)\.\s+(.+)'), # 1.1.1. Sub-subtítulo
        re.compile(r'^Cláusula\s+(\d+[a-z]?)\s*[-–]\s*(.+)', re.IGNORECASE),
        re.compile(r'^CLÁUSULA\s+(\d+[a-z]?)\s*[-–]\s*(.+)', re.IGNORECASE),
        re.compile(r'^([A-Z\s]{3,})\s*$'),  # TÍTULOS EM MAIÚSCULAS
        re.compile(r'^Seção\s+(\d+[a-z]?)\s*[-–]\s*(.+)', re.IGNORECASE),
    ]
    
    clausula_atual = None
    conteudo_linhas = []
    ultimo_numero = None
    
    for item in estrutura:
        texto = item['texto']
        encontrou_clausula = False
        
        # Tenta cada padrão
        for padrao in padroes:
            match = padrao.match(texto)
            if match:
                # Salva a cláusula anterior
                if clausula_atual:
                    clausulas.append({
                        'numero': clausula_atual['numero'],
                        'titulo': clausula_atual['titulo'],
                        'conteudo': '\n'.join(conteudo_linhas),
                        'tipo': clausula_atual.get('tipo', 'clausula'),
                        'nivel': clausula_atual.get('nivel')
                    })
                    conteudo_linhas = []
                
                # Nova cláusula
                if len(match.groups()) >= 2:
                    numero = match.group(1)
                    titulo = match.group(2) if len(match.groups()) >= 2 else texto
                else:
                    numero = f"TITULO_{len(clausulas)+1}"
                    titulo = match.group(1) if match.groups() else texto
                
                clausula_atual = {
                    'numero': numero,
                    'titulo': titulo.strip(),
                    'tipo': 'titulo' if texto.isupper() else 'clausula',
                    'nivel': item.get('nivel')
                }
                ultimo_numero = numero
                encontrou_clausula = True
                break
        
        if not encontrou_clausula and clausula_atual:
            # Adiciona ao conteúdo da cláusula atual
            if texto and texto != clausula_atual['titulo']:
                conteudo_linhas.append(texto)
        elif not encontrou_clausula and not clausula_atual:
            # Conteúdo antes da primeira cláusula
            if texto.strip():
                conteudo_linhas.append(texto)
    
    # Adiciona a última cláusula
    if clausula_atual:
        clausulas.append({
            'numero': clausula_atual['numero'],
            'titulo': clausula_atual['titulo'],
            'conteudo': '\n'.join(conteudo_linhas),
            'tipo': clausula_atual.get('tipo', 'clausula'),
            'nivel': clausula_atual.get('nivel')
        })
    
    return clausulas

def gerar_catalogo_sugestoes(clausulas, nome_doc):
    """Gera catálogo com sugestões de explicação para cada cláusula"""
    catalogo = {
        'documento': nome_doc,
        'total_clausulas': len(clausulas),
        'data_analise': '2025-10-28',
        'clausulas_com_sugestoes': []
    }
    
    # Dicionário de sugestões baseado no tipo de cláusula
    sugestoes_padrao = {
        'DEFINIÇÕES': 'Esta cláusula estabelece os termos e definições utilizados ao longo do documento, garantindo interpretação uniforme.',
        'APROVAÇÃO': 'Define os procedimentos e requisitos para aprovação da emissão dos certificados.',
        'OBJETO': 'Descreve o objeto principal do termo e os créditos imobiliários que lastreiam os CRI.',
        'IDENTIFICAÇÃO': 'Identifica os CRI emitidos, suas características e forma de distribuição.',
        'SUBSCRIÇÃO': 'Estabelece as regras para subscrição e integralização dos certificados.',
        'REMUNERAÇÃO': 'Define a forma de remuneração, atualização monetária e cálculo de juros dos CRI.',
        'VENCIMENTO': 'Estabelece as hipóteses e procedimentos de vencimento antecipado.',
        'ASSEMBLEIA': 'Regula a convocação e funcionamento das assembleias de titulares de CRI.',
        'OBRIGAÇÕES': 'Lista as obrigações e declarações da emissora perante os titulares.',
        'RESGATE': 'Define as condições e procedimentos para resgate antecipado dos CRI.',
        'AGENTE FIDUCIÁRIO': 'Estabelece as atribuições, responsabilidades e remuneração do agente fiduciário.',
        'AMORTIZAÇÃO': 'Define o cronograma e procedimentos de amortização dos certificados.',
        'GARANTIAS': 'Descreve as garantias constituídas em favor dos titulares dos CRI.',
        'EVENTOS': 'Lista os eventos de vencimento antecipado e suas consequências.',
        'COMUNICAÇÕES': 'Estabelece os procedimentos para comunicações entre as partes.',
        'FORO': 'Define o foro competente para dirimir conflitos relacionados ao documento.',
    }
    
    for clausula in clausulas:
        # Gera sugestão de explicação
        titulo_upper = clausula['titulo'].upper()
        sugestao = None
        
        # Busca por palavras-chave no título
        for palavra_chave, explicacao in sugestoes_padrao.items():
            if palavra_chave in titulo_upper:
                sugestao = explicacao
                break
        
        # Sugestão genérica se não encontrar específica
        if not sugestao:
            if clausula.get('tipo') == 'titulo':
                sugestao = f"Seção que trata de {clausula['titulo'].lower()}."
            else:
                sugestao = f"Cláusula que regula aspectos relacionados a {clausula['titulo'].lower()}."
        
        # Análise de conteúdo para melhorar sugestão
        conteudo = clausula['conteudo'].lower()
        if len(conteudo) > 100:
            if 'deverá' in conteudo or 'obriga' in conteudo:
                sugestao += ' Estabelece obrigações específicas.'
            if 'poderá' in conteudo or 'faculdade' in conteudo:
                sugestao += ' Confere faculdades e direitos.'
            if 'vedado' in conteudo or 'proibido' in conteudo:
                sugestao += ' Impõe restrições e limitações.'
            if 'prazo' in conteudo:
                sugestao += ' Especifica prazos importantes.'
            if 'comunicar' in conteudo or 'notificar' in conteudo:
                sugestao += ' Requer comunicações obrigatórias.'
        
        catalogo['clausulas_com_sugestoes'].append({
            'numero': clausula['numero'],
            'titulo': clausula['titulo'],
            'tipo': clausula.get('tipo', 'clausula'),
            'nivel': clausula.get('nivel'),
            'preview_conteudo': clausula['conteudo'][:200] + '...' if len(clausula['conteudo']) > 200 else clausula['conteudo'],
            'tamanho_conteudo': len(clausula['conteudo']),
            'sugestao_explicacao': sugestao,
            'pontos_atencao': extrair_pontos_atencao(clausula['conteudo'])
        })
    
    return catalogo

def extrair_pontos_atencao(conteudo):
    """Extrai pontos de atenção do conteúdo da cláusula"""
    pontos = []
    conteudo_lower = conteudo.lower()
    
    palavras_chave = {
        'prazo': 'Contém especificação de prazos',
        'vencimento': 'Define condições de vencimento',
        'garantia': 'Menciona garantias',
        'penalidade': 'Estabelece penalidades',
        'multa': 'Prevê aplicação de multas',
        'juros': 'Especifica taxas de juros',
        'mora': 'Regula situações de mora',
        'inadimplemento': 'Trata de inadimplemento',
        'rescisão': 'Prevê hipóteses de rescisão',
        'assembleia': 'Requer aprovação em assembleia',
        'maioria': 'Define quórum necessário',
        'notificação': 'Exige notificação prévia',
        'comunicação': 'Requer comunicação às partes',
        'registro': 'Necessita registro',
        'publicação': 'Requer publicação',
    }
    
    for palavra, descricao in palavras_chave.items():
        if palavra in conteudo_lower:
            pontos.append(descricao)
    
    return pontos if pontos else ['Revisar conteúdo completo da cláusula']

def main():
    """Função principal"""
    print("="*80)
    print("EXTRAÇÃO COMPLETA DE CLÁUSULAS - DOCUMENTOS GOLD")
    print("="*80)
    
    documentos = [
        ('GOLD_CRI_DESTINACAO_MODELO_2025.docx', 'CRI com Destinação - Modelo 2025'),
        ('GOLD_CRI_SEM_DESTINACAO_PADRAO.docx', 'CRI sem Destinação - Padrão')
    ]
    
    for doc, nome in documentos:
        print(f"\n\n{'='*80}")
        print(f"📄 Processando: {nome}")
        print(f"{'='*80}\n")
        
        caminho = f'/home/user/webapp/documentos_gold/{doc}'
        if not Path(caminho).exists():
            print(f"⚠ Arquivo não encontrado: {caminho}")
            continue
        
        # Extrai estrutura
        print("1. Extraindo estrutura do documento...")
        estrutura = extrair_estrutura_completa(caminho)
        print(f"   ✓ {len(estrutura)} elementos extraídos")
        
        # Processa cláusulas
        print("2. Identificando cláusulas...")
        clausulas = processar_clausulas_detalhadas(estrutura)
        print(f"   ✓ {len(clausulas)} cláusulas identificadas")
        
        # Gera catálogo com sugestões
        print("3. Gerando catálogo com sugestões de explicação...")
        catalogo = gerar_catalogo_sugestoes(clausulas, nome)
        
        # Salva arquivos
        base_nome = Path(caminho).stem
        
        # Arquivo completo JSON
        arquivo_json = f'/home/user/webapp/documentos_gold/{base_nome}_catalogo_completo.json'
        with open(arquivo_json, 'w', encoding='utf-8') as f:
            json.dump(catalogo, f, ensure_ascii=False, indent=2)
        print(f"   ✓ Catálogo salvo: {base_nome}_catalogo_completo.json")
        
        # Arquivo resumido markdown
        arquivo_md = f'/home/user/webapp/documentos_gold/{base_nome}_CATALOGO.md'
        with open(arquivo_md, 'w', encoding='utf-8') as f:
            f.write(f"# Catálogo de Cláusulas - {nome}\n\n")
            f.write(f"**Data de Análise:** 2025-10-28\n\n")
            f.write(f"**Total de Cláusulas:** {len(clausulas)}\n\n")
            f.write("---\n\n")
            
            for i, item in enumerate(catalogo['clausulas_com_sugestoes'], 1):
                f.write(f"## {i}. Cláusula {item['numero']}\n\n")
                f.write(f"**Título:** {item['titulo']}\n\n")
                f.write(f"**Tipo:** {item['tipo']}\n\n")
                f.write(f"### 📝 Sugestão de Explicação\n\n")
                f.write(f"{item['sugestao_explicacao']}\n\n")
                
                if item['pontos_atencao']:
                    f.write(f"### ⚠️ Pontos de Atenção\n\n")
                    for ponto in item['pontos_atencao']:
                        f.write(f"- {ponto}\n")
                    f.write("\n")
                
                f.write(f"### 📄 Preview do Conteúdo\n\n")
                f.write(f"```\n{item['preview_conteudo']}\n```\n\n")
                f.write(f"**Tamanho do conteúdo:** {item['tamanho_conteudo']} caracteres\n\n")
                f.write("---\n\n")
        
        print(f"   ✓ Catálogo markdown salvo: {base_nome}_CATALOGO.md")
        
        # Mostra resumo
        print(f"\n📊 RESUMO:")
        print(f"   • Total de cláusulas: {len(clausulas)}")
        print(f"   • Cláusulas principais: {sum(1 for c in clausulas if c.get('tipo') == 'clausula')}")
        print(f"   • Títulos/Seções: {sum(1 for c in clausulas if c.get('tipo') == 'titulo')}")
    
    print("\n\n" + "="*80)
    print("✅ ANÁLISE CONCLUÍDA COM SUCESSO!")
    print("="*80)
    print("\n📁 Arquivos gerados:")
    print("   • *_catalogo_completo.json - Dados completos em JSON")
    print("   • *_CATALOGO.md - Catálogo formatado em Markdown")
    print("\nPróximos passos:")
    print("   1. Revisar os catálogos gerados")
    print("   2. Ajustar sugestões de explicação conforme necessário")
    print("   3. Integrar com sistema de revisão de documentos")

if __name__ == '__main__':
    main()
