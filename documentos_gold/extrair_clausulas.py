#!/usr/bin/env python3
"""
Script para extrair e analisar cláusulas dos documentos GOLD
"""
import json
from docx import Document
from pathlib import Path
import re

def extrair_texto_documento(caminho_docx):
    """Extrai todo o texto de um documento DOCX"""
    doc = Document(caminho_docx)
    texto_completo = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            texto_completo.append(para.text.strip())
    
    # Também extrair texto de tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texto_completo.append(cell.text.strip())
    
    return "\n".join(texto_completo)

def identificar_clausulas(texto):
    """Identifica cláusulas e suas seções no documento"""
    clausulas = []
    linhas = texto.split('\n')
    
    clausula_atual = None
    conteudo_atual = []
    
    # Padrões para identificar cláusulas
    padrao_clausula = re.compile(r'^(\d+\.|\d+\.\d+\.?)\s+(.+)', re.IGNORECASE)
    padrao_secao = re.compile(r'^(SEÇÃO|CAPÍTULO|TÍTULO)\s+(.+)', re.IGNORECASE)
    
    for i, linha in enumerate(linhas):
        # Verifica se é uma seção
        match_secao = padrao_secao.match(linha)
        if match_secao:
            if clausula_atual:
                clausulas.append({
                    'numero': clausula_atual,
                    'titulo': titulo_clausula,
                    'conteudo': '\n'.join(conteudo_atual)
                })
            clausula_atual = f"SECAO_{match_secao.group(2)}"
            titulo_clausula = match_secao.group(2)
            conteudo_atual = []
            continue
        
        # Verifica se é uma cláusula numerada
        match_clausula = padrao_clausula.match(linha)
        if match_clausula:
            # Salva a cláusula anterior se existir
            if clausula_atual:
                clausulas.append({
                    'numero': clausula_atual,
                    'titulo': titulo_clausula,
                    'conteudo': '\n'.join(conteudo_atual)
                })
            
            clausula_atual = match_clausula.group(1)
            titulo_clausula = match_clausula.group(2)
            conteudo_atual = []
        elif clausula_atual:
            # Adiciona conteúdo à cláusula atual
            if linha.strip():
                conteudo_atual.append(linha)
    
    # Adiciona a última cláusula
    if clausula_atual:
        clausulas.append({
            'numero': clausula_atual,
            'titulo': titulo_clausula,
            'conteudo': '\n'.join(conteudo_atual)
        })
    
    return clausulas

def analisar_documento(caminho_docx, nome_doc):
    """Analisa um documento e extrai suas cláusulas"""
    print(f"\n{'='*80}")
    print(f"Analisando: {nome_doc}")
    print(f"{'='*80}\n")
    
    texto = extrair_texto_documento(caminho_docx)
    clausulas = identificar_clausulas(texto)
    
    print(f"Total de cláusulas identificadas: {len(clausulas)}\n")
    
    # Salvar análise em arquivo JSON
    resultado = {
        'documento': nome_doc,
        'total_clausulas': len(clausulas),
        'clausulas': clausulas
    }
    
    # Salvar em arquivo
    nome_saida = Path(caminho_docx).stem + '_analise.json'
    with open(f'/home/user/webapp/documentos_gold/{nome_saida}', 'w', encoding='utf-8') as f:
        json.dump(resultado, f, ensure_ascii=False, indent=2)
    
    print(f"✓ Análise salva em: {nome_saida}")
    
    # Mostrar primeiras 10 cláusulas
    print("\nPrimeiras cláusulas identificadas:")
    print("-" * 80)
    for i, clausula in enumerate(clausulas[:10], 1):
        print(f"\n{i}. Cláusula {clausula['numero']}")
        print(f"   Título: {clausula['titulo']}")
        conteudo_preview = clausula['conteudo'][:150] + '...' if len(clausula['conteudo']) > 150 else clausula['conteudo']
        print(f"   Conteúdo: {conteudo_preview}")
    
    return resultado

def main():
    """Função principal"""
    documentos = [
        ('GOLD_CRI_DESTINACAO_MODELO_2025.docx', 'CRI com Destinação'),
        ('GOLD_CRI_SEM_DESTINACAO_PADRAO.docx', 'CRI sem Destinação')
    ]
    
    resultados = []
    
    for doc, nome in documentos:
        caminho = f'/home/user/webapp/documentos_gold/{doc}'
        if Path(caminho).exists():
            resultado = analisar_documento(caminho, nome)
            resultados.append(resultado)
        else:
            print(f"⚠ Documento não encontrado: {caminho}")
    
    # Criar resumo geral
    print(f"\n{'='*80}")
    print("RESUMO GERAL")
    print(f"{'='*80}\n")
    
    for resultado in resultados:
        print(f"📄 {resultado['documento']}")
        print(f"   Total de cláusulas: {resultado['total_clausulas']}")
    
    print("\n✓ Análise concluída!")
    print("✓ Arquivos JSON criados na pasta documentos_gold/")

if __name__ == '__main__':
    main()
