from typing import List, Dict
from docx import Document

# estilos comuns de heading em PT/EN
HEADING_STYLES = {f'Heading {i}' for i in range(1,10)} | {'Título 1','Título 2','Título 3','Título 4','Título 5'}

def extract_blocks_docx(path: str) -> List[Dict]:
    """
    Extrai blocos de texto de um arquivo DOCX preservando estrutura hierárquica.

    Retorna blocos com section_path, para_idx, table_ref e text.

    Args:
        path: Caminho para o arquivo DOCX

    Returns:
        Lista de dicionários com {mode, section_path, para_idx, table_ref, text}
    """
    doc = Document(path)
    section_stack = []
    blocks: List[Dict] = []
    para_counter = 0

    def section_path() -> str:
        return " > ".join(section_stack) if section_stack else ""

    # percorre parágrafos em ordem
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or ""
        text = (p.text or "").strip()
        if not text:
            continue
        if style in HEADING_STYLES:
            # abre nova seção
            section_stack.append(text)
            para_counter = 0
            continue
        para_counter += 1
        blocks.append({
            "mode": "docx",
            "section_path": section_path(),
            "para_idx": para_counter,
            "text": text
        })

    # percorre tabelas (se houver)
    # obs: python-docx não preserva posição relativa perfeita; serve como evidência de conteúdo
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                t = (cell.text or "").strip()
                if not t:
                    continue
                blocks.append({
                    "mode": "docx",
                    "section_path": section_path(),
                    "table_ref": f"Tabela {ti+1} / R{ri+1}C{ci+1}",
                    "text": t
                })

    return blocks
