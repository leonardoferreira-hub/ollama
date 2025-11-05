"""
Módulo para geração de relatórios em Excel e DOCX
"""

from typing import List, Dict
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment


def generate_excel_report(reviews: List[Dict], output_path: Path):
    """
    Gera relatório em Excel com análise das cláusulas

    Args:
        reviews: Lista de revisões
        output_path: Caminho para salvar o Excel
    """
    data = []

    for review in reviews:
        clause = review['clause']
        references = review['references']

        # Pega melhor match
        best_match = references[0] if references else None
        best_score = best_match['combined_score'] if best_match else 0

        # Extrai prioridade da revisão (parsing básico)
        review_text = review['review']
        priority = 'Média'
        if 'PRIORIDADE: Alta' in review_text or 'Prioridade: Alta' in review_text:
            priority = 'Alta'
        elif 'PRIORIDADE: Baixa' in review_text or 'Prioridade: Baixa' in review_text:
            priority = 'Baixa'

        data.append({
            'Cláusula': clause['title'],
            'Seção': clause.get('section', ''),
            'Score Similaridade': f"{best_score:.2%}",
            'Prioridade': priority,
            'Categoria Referência': best_match['catalog_clause']['category'] if best_match else '',
            'Importância': best_match['catalog_clause']['importance'] if best_match else '',
            'Conteúdo Original': clause['content'][:500] + '...' if len(clause['content']) > 500 else clause['content'],
            'Análise': review_text
        })

    # Cria DataFrame
    df = pd.DataFrame(data)

    # Salva com formatação
    excel_path = output_path / 'revisao.xlsx'
    df.to_excel(excel_path, index=False, sheet_name='Revisão')

    # Formatação adicional
    wb = openpyxl.load_workbook(excel_path)
    ws = wb['Revisão']

    # Header
    header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # Ajusta larguras
    ws.column_dimensions['A'].width = 40
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 25
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 60
    ws.column_dimensions['H'].width = 80

    # Cores por prioridade
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        priority_cell = row[3]  # Coluna D (Prioridade)
        if priority_cell.value == 'Alta':
            for cell in row:
                cell.fill = PatternFill(start_color='FFE6E6', end_color='FFE6E6', fill_type='solid')
        elif priority_cell.value == 'Baixa':
            for cell in row:
                cell.fill = PatternFill(start_color='E6F7E6', end_color='E6F7E6', fill_type='solid')

    wb.save(excel_path)


def generate_docx_report(reviews: List[Dict], document_info, output_path: Path):
    """
    Gera relatório narrativo em DOCX

    Args:
        reviews: Lista de revisões
        document_info: Informações do documento original
        output_path: Caminho para salvar o DOCX
    """
    doc = Document()

    # Título
    title = doc.add_heading('Relatório de Revisão Jurídica', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadados
    doc.add_paragraph(f'Documento: {document_info.filename}')
    doc.add_paragraph(f'Total de Cláusulas: {len(reviews)}')
    doc.add_paragraph('')

    # Sumário executivo
    doc.add_heading('Sumário Executivo', 1)

    high_priority = sum(1 for r in reviews if 'Alta' in r['review'])
    doc.add_paragraph(f'Cláusulas com prioridade ALTA: {high_priority}')

    if high_priority > 0:
        p = doc.add_paragraph('⚠️ ', style='List Bullet')
        run = p.add_run('Atenção: Existem cláusulas que requerem revisão urgente.')
        run.bold = True

    doc.add_page_break()

    # Análise detalhada
    doc.add_heading('Análise Detalhada', 1)

    for i, review in enumerate(reviews, 1):
        clause = review['clause']

        # Título da cláusula
        heading = doc.add_heading(f"{i}. {clause['title']}", 2)

        # Info básica
        if clause.get('section'):
            doc.add_paragraph(f"Seção: {clause['section']}", style='List Bullet')

        # Melhor match
        if review['references']:
            best = review['references'][0]
            score = best['combined_score']
            p = doc.add_paragraph(
                f"Similaridade com catálogo: {score:.1%} - ",
                style='List Bullet'
            )
            run = p.add_run(best['catalog_clause']['title'])
            run.italic = True

        # Conteúdo original (resumido)
        doc.add_heading('Texto Original:', 3)
        content = clause['content']
        if len(content) > 300:
            content = content[:300] + '...'
        doc.add_paragraph(content)

        # Análise
        doc.add_heading('Análise:', 3)
        analysis_para = doc.add_paragraph(review['review'])

        # Separador
        doc.add_paragraph('_' * 80)

    # Rodapé
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.add_run('Relatório gerado automaticamente por Jurídico Review AI\n').italic = True
    footer.add_run(f'Modelo: {review["model"]} ({review["provider"]})').italic = True

    # Salva
    docx_path = output_path / 'sugestoes.docx'
    doc.save(docx_path)


def generate_reports(reviews: List[Dict], document_info, output_path: Path):
    """
    Gera ambos os relatórios (Excel e DOCX)

    Args:
        reviews: Lista de revisões
        document_info: Informações do documento
        output_path: Diretório de saída
    """
    generate_excel_report(reviews, output_path)
    generate_docx_report(reviews, document_info, output_path)
