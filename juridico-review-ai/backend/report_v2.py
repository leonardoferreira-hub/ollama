"""
Geração de Relatórios v2: Excel e DOCX com análise completa
Tier-1 + Tier-2 + Auditoria
"""

from typing import List, Dict
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


def generate_excel_report(tier1_results: List[Dict],
                          tier2_results: List[Dict],
                          output_path: Path,
                          timestamp: str = None):
    """
    Gera Excel com múltiplas abas

    Abas:
    1. Resumo Executivo
    2. Cláusulas OK (Tier-1)
    3. Sugestões (Tier-2)
    4. Detalhamento Completo
    """

    if not timestamp:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    excel_file = output_path / f'revisao_completa_{timestamp}.xlsx'

    with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
        # ====== ABA 1: RESUMO EXECUTIVO ======
        resumo_data = {
            'Métrica': [
                'Total de Cláusulas Analisadas',
                'Cláusulas OK (Tier-1)',
                'Cláusulas com Sugestões (Tier-2)',
                'PRESENTES',
                'PARCIAIS',
                'AUSENTES',
                'Cláusulas Obrigatórias Atendidas',
                'Cláusulas Obrigatórias Pendentes'
            ],
            'Valor': [
                len(tier1_results) + len(tier2_results),
                len(tier1_results),
                len(tier2_results),
                sum(1 for r in tier1_results
                   if r['classification']['classificacao'] == 'PRESENTE'),
                sum(1 for r in tier2_results
                   if r['classification']['classificacao'] == 'PARCIAL'),
                sum(1 for r in tier2_results
                   if r['classification']['classificacao'] == 'AUSENTE'),
                sum(1 for r in tier1_results
                   if r['classification'].get('obrigatoria')),
                sum(1 for r in tier2_results
                   if r['classification'].get('obrigatoria'))
            ]
        }

        df_resumo = pd.DataFrame(resumo_data)
        df_resumo.to_excel(writer, sheet_name='Resumo', index=False)

        # ====== ABA 2: TIER-1 (OK) ======
        tier1_data = []
        for item in tier1_results:
            clause = item['clause']
            classif = item['classification']

            tier1_data.append({
                'Cláusula': clause['title'],
                'Seção': clause.get('section', ''),
                'Status': classif['classificacao'],
                'Confiança': f"{classif.get('confianca', 0):.0%}",
                'Catálogo ID': classif.get('catalog_id', ''),
                'Catálogo Título': classif.get('catalog_titulo', ''),
                'Obrigatória': 'Sim' if classif.get('obrigatoria') else 'Não',
                'Justificativa': classif.get('justificativa', '')
            })

        if tier1_data:
            df_tier1 = pd.DataFrame(tier1_data)
            df_tier1.to_excel(writer, sheet_name='Cláusulas OK', index=False)

        # ====== ABA 3: TIER-2 (SUGESTÕES) ======
        tier2_data = []
        for item in tier2_results:
            if item.get('skipped'):
                continue

            clause = item['clause']
            classif = item['classification']
            suggestion = item.get('suggestion', {})

            tier2_data.append({
                'Cláusula': clause['title'],
                'Status Original': classif['classificacao'],
                'Obrigatória': 'Sim' if classif.get('obrigatoria') else 'Não',
                'Catálogo ID': suggestion.get('catalog_id', ''),
                'Tipo Sugestão': suggestion.get('sugestao_tipo', ''),
                'Prioridade': suggestion.get('nivel_prioridade', 'MEDIA'),
                'Conformidade Estimada': f"{suggestion.get('conformidade_estimada', 0):.0%}",
                'Explicação': suggestion.get('explicacao', ''),
                'Texto Sugerido': suggestion.get('texto_sugerido', '')[:500] + '...'
                                 if len(suggestion.get('texto_sugerido', '')) > 500
                                 else suggestion.get('texto_sugerido', '')
            })

        if tier2_data:
            df_tier2 = pd.DataFrame(tier2_data)
            df_tier2.to_excel(writer, sheet_name='Sugestões', index=False)

        # ====== ABA 4: DETALHAMENTO COMPLETO ======
        all_data = []

        for item in tier1_results:
            all_data.append(_format_full_row(item, tier='1'))

        for item in tier2_results:
            if not item.get('skipped'):
                all_data.append(_format_full_row(item, tier='2'))

        if all_data:
            df_all = pd.DataFrame(all_data)
            df_all.to_excel(writer, sheet_name='Detalhamento', index=False)

    # Formatação
    _format_excel(excel_file)

    logger.info(f"Excel gerado: {excel_file}")
    return excel_file


def _format_full_row(item: Dict, tier: str) -> Dict:
    """Formata linha completa para aba de detalhamento"""
    clause = item['clause']
    classif = item['classification']
    suggestion = item.get('suggestion', {})

    row = {
        'Tier': tier,
        'Cláusula': clause['title'],
        'Seção': clause.get('section', ''),
        'Status': classif['classificacao'],
        'Confiança': classif.get('confianca', 0),
        'Obrigatória': 'Sim' if classif.get('obrigatoria') else 'Não',
        'Catálogo ID': classif.get('catalog_id', ''),
        'Match Score': classif.get('match_score', 0),
    }

    if tier == '2' and suggestion:
        row['Tipo Sugestão'] = suggestion.get('sugestao_tipo', '')
        row['Prioridade'] = suggestion.get('nivel_prioridade', '')
        row['Provider'] = suggestion.get('provider', '')
        row['Model'] = suggestion.get('model', '')

    return row


def _format_excel(filepath: Path):
    """Aplica formatação ao Excel"""
    wb = openpyxl.load_workbook(filepath)

    header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # Header
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border

        # Auto-width
        for column_cells in ws.columns:
            length = max(len(str(cell.value or '')) for cell in column_cells)
            ws.column_dimensions[column_cells[0].column_letter].width = min(length + 2, 80)

    wb.save(filepath)


def generate_docx_report(tier1_results: List[Dict],
                        tier2_results: List[Dict],
                        document_info,
                        catalog_info: Dict,
                        output_path: Path,
                        timestamp: str = None):
    """
    Gera DOCX narrativo com sugestões detalhadas
    """
    doc = Document()

    # ====== TÍTULO ======
    title = doc.add_heading('Relatório de Revisão de Minuta', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ====== METADADOS ======
    doc.add_paragraph(f'Documento: {document_info.filename}')
    doc.add_paragraph(f'Catálogo: {catalog_info.get("metadata", {}).get("nome")} '
                     f'v{catalog_info.get("metadata", {}).get("versao")}')
    doc.add_paragraph('')

    # ====== SUMÁRIO EXECUTIVO ======
    doc.add_heading('Sumário Executivo', 1)

    total = len(tier1_results) + len(tier2_results)
    doc.add_paragraph(f'Total de cláusulas analisadas: {total}')
    doc.add_paragraph(f'✓ Aprovadas (Tier-1): {len(tier1_results)}', style='List Bullet')
    doc.add_paragraph(f'⚠ Com sugestões (Tier-2): {len(tier2_results)}', style='List Bullet')

    if tier2_results:
        ausentes = sum(1 for r in tier2_results
                      if r['classification']['classificacao'] == 'AUSENTE')
        parciais = sum(1 for r in tier2_results
                      if r['classification']['classificacao'] == 'PARCIAL')

        if ausentes > 0:
            p = doc.add_paragraph()
            run = p.add_run(f'❌ ATENÇÃO: {ausentes} cláusula(s) AUSENTE(S)')
            run.bold = True
            run.font.color.rgb = RGBColor(200, 0, 0)

        if parciais > 0:
            p = doc.add_paragraph()
            run = p.add_run(f'⚠ {parciais} cláusula(s) PARCIAL(IS) - requer complementação')
            run.font.color.rgb = RGBColor(255, 140, 0)

    doc.add_page_break()

    # ====== SUGESTÕES TIER-2 ======
    if tier2_results:
        doc.add_heading('Sugestões e Correções', 1)

        for i, item in enumerate(tier2_results, 1):
            if item.get('skipped'):
                continue

            clause = item['clause']
            classif = item['classification']
            suggestion = item.get('suggestion', {})

            # Título da cláusula
            heading = doc.add_heading(f"{i}. {clause['title']}", 2)

            # Status
            status = classif['classificacao']
            p = doc.add_paragraph()
            run = p.add_run(f"Status: {status}")
            run.bold = True

            if status == 'AUSENTE':
                run.font.color.rgb = RGBColor(200, 0, 0)
            elif status == 'PARCIAL':
                run.font.color.rgb = RGBColor(255, 140, 0)

            # Prioridade
            prioridade = suggestion.get('nivel_prioridade', 'MEDIA')
            doc.add_paragraph(f"Prioridade: {prioridade}", style='List Bullet')
            doc.add_paragraph(f"Tipo: {suggestion.get('sugestao_tipo', '')}", style='List Bullet')

            # Explicação
            if suggestion.get('explicacao'):
                doc.add_heading('Análise:', 3)
                doc.add_paragraph(suggestion['explicacao'])

            # Texto sugerido
            if suggestion.get('texto_sugerido'):
                doc.add_heading('Texto Sugerido:', 3)
                p = doc.add_paragraph(suggestion['texto_sugerido'])
                p.style = 'Intense Quote'

            # Checklist
            if suggestion.get('itens_checklist'):
                doc.add_heading('Checklist de Validação:', 3)
                for item_check in suggestion['itens_checklist']:
                    doc.add_paragraph(f'☐ {item_check}', style='List Bullet')

            doc.add_paragraph('_' * 80)

    # ====== CLÁUSULAS OK ======
    if tier1_results:
        doc.add_page_break()
        doc.add_heading('Cláusulas Aprovadas (Tier-1)', 1)
        doc.add_paragraph(f'{len(tier1_results)} cláusula(s) atenderam os critérios:')

        for item in tier1_results:
            clause = item['clause']
            classif = item['classification']

            doc.add_paragraph(
                f"✓ {clause['title']} "
                f"(Confiança: {classif.get('confianca', 0):.0%})",
                style='List Bullet'
            )

    # ====== RODAPÉ ======
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.add_run('Relatório gerado por Jurídico Review AI v2.0\n').italic = True
    footer.add_run('Sistema de revisão automatizada com arquitetura em camadas').italic = True

    # Salva
    if not timestamp:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    docx_path = output_path / f'sugestoes_detalhadas_{timestamp}.docx'
    doc.save(docx_path)

    logger.info(f"DOCX gerado: {docx_path}")
    return docx_path


def generate_comprehensive_reports(tier1_results: List[Dict],
                                   tier2_results: List[Dict],
                                   document_info,
                                   catalog_info: Dict,
                                   output_path: Path,
                                   audit_trail=None):
    """
    Gera todos os relatórios (Excel + DOCX)

    Args:
        tier1_results: Resultados Tier-1 (OK)
        tier2_results: Resultados Tier-2 (com sugestões)
        document_info: Info do documento
        catalog_info: Info do catálogo
        output_path: Diretório de saída
        audit_trail: Objeto de auditoria
    """
    # Unified timestamp to avoid name collisions and permission issues
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    excel_path = generate_excel_report(tier1_results, tier2_results, output_path, timestamp=timestamp)
    docx_path = generate_docx_report(tier1_results, tier2_results, document_info, catalog_info, output_path, timestamp=timestamp)

    if audit_trail:
        audit_trail.log_event('REPORTS_GENERATED', {
            'excel': excel_path.name,
            'docx': docx_path.name
        })

    return excel_path, docx_path
