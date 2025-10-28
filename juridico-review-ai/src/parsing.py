"""
Módulo para leitura e segmentação de documentos DOCX e PDF
Suporta: CLÁUSULA, CAPÍTULO, ANEXO, tabelas e placeholders
"""

from pathlib import Path
from typing import List, Dict
import docx
import PyPDF2
import pdfplumber
import re


class Document:
    """Representa um documento parseado"""

    def __init__(self, filepath: str):
        self.filepath = Path(filepath)
        self.filename = self.filepath.name
        self.clauses = []
        self.metadata = {}

    def add_clause(self, title: str, content: str, section: str = None, source: str = "paragraph"):
        """Adiciona uma cláusula ao documento"""
        self.clauses.append({
            'title': title,
            'content': content,
            'section': section,
            'source': source,  # paragraph, table, anexo, capitulo
            'index': len(self.clauses)
        })


def normalize_text(text: str) -> str:
    """
    Normaliza texto removendo placeholders e field codes
    Remove: {=}, [=], etc.
    """
    # Remove placeholders
    text = re.sub(r'\{=\}|\[=\]|\{.*?\}', '', text)

    # Remove múltiplos espaços
    text = re.sub(r'\s+', ' ', text)

    return text.strip()


def parse_document(filepath: str) -> Document:
    """
    Parse de documento DOCX ou PDF

    Args:
        filepath: Caminho para o arquivo

    Returns:
        Objeto Document com cláusulas extraídas
    """
    path = Path(filepath)
    doc = Document(filepath)

    if path.suffix.lower() == '.docx':
        return parse_docx(doc)
    elif path.suffix.lower() == '.pdf':
        return parse_pdf(doc)
    else:
        raise ValueError(f"Formato não suportado: {path.suffix}")


def extract_table_content(table) -> List[Dict]:
    """
    Extrai conteúdo de tabelas do DOCX
    Retorna lista de células relevantes
    """
    table_data = []

    for i, row in enumerate(table.rows):
        row_text = []
        for cell in row.cells:
            cell_text = normalize_text(cell.text)
            if cell_text:
                row_text.append(cell_text)

        if row_text:
            # Junta células da linha
            full_row = ' | '.join(row_text)
            table_data.append({
                'row_index': i,
                'text': full_row
            })

    return table_data


def parse_docx(doc: Document) -> Document:
    """
    Parse ULTRA ROBUSTO para arquivos DOCX

    Captura múltiplos padrões:
    - CLÁUSULA X - TÍTULO
    - CAPÍTULO NOME  
    - ANEXO NOME
    - Numeração: 1., 1.1, 1.1.1
    - Títulos com formatação (Bold, Heading)
    - Texto em MAIÚSCULAS como título
    - Conteúdo de TABELAS
    """
    docx_file = docx.Document(doc.filepath)

    current_clause = None
    current_content = []
    current_section = None

    # Processa PARÁGRAFOS
    for idx, para in enumerate(docx_file.paragraphs):
        text = normalize_text(para.text)

        if not text or len(text) < 3:
            continue

        # ====== DETECÇÃO DE TÍTULOS/CLÁUSULAS ======
        
        # 1. REGEX para palavras-chave comuns
        heading_match = re.match(
            r'^(cl[áa]usula|se[çc][ãa]o|cap[íi]tulo|anexo|item|art|artigo|t[íi]tulo|par[áa]grafo)\s+(.+)$',
            text,
            re.IGNORECASE
        )

        # 2. Numeração com ou sem título (muito flexível)
        number_match = re.match(
            r'^(\d+(\.\d+){0,3})[\.\)\s]*[-–—]?\s*(.*)$',
            text
        )

        # 3. Letras (a), (b), etc.
        letter_match = re.match(
            r'^([a-z]\)|[a-z]\.|[A-Z]\)|[A-Z]\.)\s+(.+)$',
            text
        )

        # 4. Verifica formatação (Bold, Heading, etc.)
        is_bold = False
        is_heading_style = False
        try:
            style_name = getattr(para.style, 'name', '').lower()
            is_heading_style = 'heading' in style_name or 'título' in style_name or 'title' in style_name
            
            # Verifica se tem runs em negrito
            for run in para.runs:
                if run.bold:
                    is_bold = True
                    break
        except:
            pass

        # 5. Texto em MAIÚSCULAS (possível título)
        is_upper = text.isupper() and len(text.split()) >= 3 and len(text.split()) <= 15

        # 6. Linha curta seguida de dois pontos (ex: "Prazo:")
        colon_match = re.match(r'^([^:]+):\s*$', text)

        # ====== DECIDE SE É TÍTULO ======
        is_title = False
        
        if heading_match:
            is_title = True
        elif number_match and number_match.group(3):  # Tem número E texto
            # É título se: tem formatação E (texto razoável OU maiúsculas)
            has_text = len(number_match.group(3).strip()) > 10
            is_short = len(text.split()) <= 20
            # Só considera título se tiver formatação OU for maiúsculas
            if has_text and (is_bold or is_heading_style or (is_upper and is_short)):
                is_title = True
        elif letter_match and len(text.split()) >= 3:  # Letras só se tiver texto razoável
            is_title = True
        elif is_heading_style:
            is_title = True
        elif is_bold and len(text.split()) >= 3 and len(text.split()) <= 15:  # Bold + comprimento razoável
            is_title = True
        elif is_upper and len(text) > 15 and len(text.split()) >= 3:  # MAIÚSCULAS com mínimo de palavras
            is_title = True
        elif colon_match and len(text) <= 80 and len(text.split()) >= 2:  # "Título:" com mínimo palavras
            is_title = True

        # ====== PROCESSA COMO TÍTULO OU CONTEÚDO ======
        if is_title:
            # Salva cláusula anterior
            if current_clause:
                doc.add_clause(
                    current_clause,
                    '\n'.join(current_content),
                    current_section,
                    source="paragraph"
                )

            # Nova cláusula
            current_clause = text
            current_content = []

            # Determina seção
            if heading_match:
                tipo = heading_match.group(1).upper()
                current_section = f"{tipo}: {heading_match.group(2)}"
            elif is_upper and len(text.split()) <= 6:
                current_section = text

        # Conteúdo da cláusula atual
        else:
            if current_clause:
                current_content.append(text)
            else:
                # Se ainda não há cláusula, cria uma genérica
                if not current_clause and len(text) > 20:
                    current_clause = f"PREÂMBULO/INTRODUÇÃO"
                    current_content.append(text)

    # Adiciona última cláusula dos parágrafos
    if current_clause:
        doc.add_clause(
            current_clause,
            '\n'.join(current_content),
            current_section,
            source="paragraph"
        )

    # ====== FALLBACK: Se encontrou poucas cláusulas, divide por parágrafos grandes ======
    min_expected_clauses = 5
    if len(doc.clauses) < min_expected_clauses:
        # Junta todos os parágrafos
        all_paragraphs = []
        for para in docx_file.paragraphs:
            text = normalize_text(para.text)
            if text and len(text) > 30:  # Ignora linhas muito curtas
                all_paragraphs.append(text)
        
        # Divide em chunks de ~500 palavras
        chunk_size = 100  # palavras por chunk
        current_chunk = []
        word_count = 0
        chunk_num = 1
        
        for para_text in all_paragraphs:
            words = para_text.split()
            word_count += len(words)
            current_chunk.append(para_text)
            
            if word_count >= chunk_size:
                # Salva chunk
                chunk_title = f"SEÇÃO {chunk_num}"
                # Tenta extrair primeira frase como título
                first_sentence = current_chunk[0][:100] if current_chunk else f"Seção {chunk_num}"
                doc.add_clause(
                    title=f"{chunk_title}: {first_sentence}",
                    content='\n\n'.join(current_chunk),
                    section=chunk_title,
                    source="fallback"
                )
                
                current_chunk = []
                word_count = 0
                chunk_num += 1
        
        # Adiciona último chunk
        if current_chunk:
            chunk_title = f"SEÇÃO {chunk_num}"
            first_sentence = current_chunk[0][:100] if current_chunk else f"Seção {chunk_num}"
            doc.add_clause(
                title=f"{chunk_title}: {first_sentence}",
                content='\n\n'.join(current_chunk),
                section=chunk_title,
                source="fallback"
            )

    # Processa TABELAS
    for table_idx, table in enumerate(docx_file.tables):
        table_content = extract_table_content(table)

        if not table_content:
            continue

        # Primeira linha geralmente é cabeçalho/título
        first_row = table_content[0]['text']

        # Junta todas as linhas
        all_rows = '\n'.join([row['text'] for row in table_content])

        # Adiciona como cláusula (com título da primeira linha)
        doc.add_clause(
            title=f"TABELA {table_idx + 1}: {first_row[:80]}",
            content=all_rows,
            section="TABELAS",
            source="table"
        )

    # Metadata
    doc.metadata['total_clauses'] = len(doc.clauses)
    doc.metadata['has_tables'] = len(docx_file.tables) > 0
    doc.metadata['num_paragraphs'] = len(docx_file.paragraphs)
    doc.metadata['num_tables'] = len(docx_file.tables)

    return doc


def parse_pdf(doc: Document) -> Document:
    """
    Parse específico para arquivos PDF
    Usa pdfplumber para melhor extração de texto e tabelas
    """
    with pdfplumber.open(doc.filepath) as pdf:
        full_text = []
        all_tables = []

        for page_num, page in enumerate(pdf.pages):
            # Extrai texto
            text = page.extract_text()
            if text:
                full_text.append(text)

            # Extrai tabelas
            tables = page.extract_tables()
            if tables:
                for table_idx, table in enumerate(tables):
                    all_tables.append({
                        'page': page_num + 1,
                        'table_idx': table_idx,
                        'data': table
                    })

    # Processa texto completo
    content = '\n'.join(full_text)
    lines = content.split('\n')

    current_clause = None
    current_content = []
    current_section = None

    for line in lines:
        text = normalize_text(line)

        if not text or len(text) < 3:
            continue

        # REGEX EXPANDIDA: CLÁUSULA | CAPÍTULO | ANEXO | SEÇÃO
        heading_match = re.match(
            r'^(cl[áa]usula|se[çc][ãa]o|cap[íi]tulo|anexo)\s+(.+)$',
            text,
            re.IGNORECASE
        )

        # Numeração
        number_match = re.match(
            r'^(\d+(\.\d+){0,2})\s*[-–—]?\s*(.+)$',
            text
        )

        if heading_match or number_match:
            if current_clause:
                doc.add_clause(
                    current_clause,
                    '\n'.join(current_content),
                    current_section,
                    source="paragraph"
                )

            current_clause = text
            current_content = []

            if heading_match:
                tipo = heading_match.group(1).upper()
                current_section = f"{tipo}: {heading_match.group(2)}"

        elif text.isupper() and len(text.split()) <= 6 and len(text) > 10:
            current_section = text

        elif current_clause:
            current_content.append(text)

    if current_clause:
        doc.add_clause(
            current_clause,
            '\n'.join(current_content),
            current_section,
            source="paragraph"
        )

    # Processa TABELAS do PDF
    for table_info in all_tables:
        table_data = table_info['data']
        page = table_info['page']

        # Junta linhas da tabela
        table_text = []
        for row in table_data:
            if row:
                row_text = ' | '.join([str(cell) if cell else '' for cell in row])
                if row_text.strip():
                    table_text.append(row_text)

        if table_text:
            first_row = table_text[0]
            all_text = '\n'.join(table_text)

            doc.add_clause(
                title=f"TABELA Pág.{page}: {first_row[:80]}",
                content=all_text,
                section="TABELAS",
                source="table"
            )

    # Metadata
    doc.metadata['total_clauses'] = len(doc.clauses)
    doc.metadata['has_tables'] = len(all_tables) > 0
    doc.metadata['num_tables'] = len(all_tables)

    return doc


def extract_paragraphs(text: str) -> List[str]:
    """Divide texto em parágrafos mantendo estrutura"""
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    return paragraphs
