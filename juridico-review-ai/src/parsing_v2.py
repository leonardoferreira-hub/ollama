"""
Parsing V2: Baseado EXCLUSIVAMENTE em estilos Heading do Word
Agrupa TODO o conteúdo entre Headings (incluindo sub-parágrafos)
"""

from pathlib import Path
from typing import List, Dict
import docx
import re


class Document:
    """Representa um documento parseado"""

    def __init__(self, filepath: str):
        self.filepath = Path(filepath)
        self.filename = self.filepath.name
        self.clauses = []
        self.metadata = {}

    def add_clause(self, title: str, content: str, level: int = 1, source: str = "heading"):
        """Adiciona uma cláusula ao documento"""
        self.clauses.append({
            'title': title,
            'content': content,
            'level': level,  # Nível hierárquico (1=principal, 2=sub, etc.)
            'source': source,
            'index': len(self.clauses)
        })


def normalize_text(text: str) -> str:
    """Remove placeholders e normaliza"""
    text = re.sub(r'\{=\}|\[=\]|\{.*?\}', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def parse_docx_by_headings(filepath: str, max_heading_level: int = 2, min_content_length: int = 50) -> Document:
    """
    Parse baseado SOMENTE nos estilos Heading do Word - VERSÃO OTIMIZADA
    
    Algoritmo INTELIGENTE:
    1. Identifica APENAS Headings de nível 1 e 2 (principais)
    2. Agrupa TODO o conteúdo subsequente (incluindo sub-headings 3, 4, 5)
    3. Ignora headings curtos sem conteúdo (títulos decorativos)
    4. Consolida headings consecutivos em um único título composto
    5. Resultado: Cláusulas principais completas com todo contexto
    
    Args:
        filepath: Caminho do arquivo DOCX
        max_heading_level: Nível máximo de heading para criar cláusula (padrão: 2)
        min_content_length: Mínimo de caracteres de conteúdo para salvar cláusula (padrão: 50)
    """
    doc = Document(filepath)
    docx_file = docx.Document(filepath)
    
    current_heading = None
    current_content = []
    current_level = 0
    pending_headings = []  # Acumula headings consecutivos sem conteúdo
    
    for para in docx_file.paragraphs:
        text = normalize_text(para.text)
        
        if not text:
            continue
        
        # Verifica se é Heading e seu nível
        is_main_heading = False
        heading_level = 0
        
        try:
            style_name = para.style.name.lower()
            
            if 'heading' in style_name or 'título' in style_name:
                # Extrai nível
                level_match = re.search(r'(\d+)', style_name)
                if level_match:
                    heading_level = int(level_match.group(1))
                else:
                    heading_level = 1
                
                # APENAS headings de nível baixo (1, 2) criam novas cláusulas
                if heading_level <= max_heading_level:
                    is_main_heading = True
        except:
            pass
        
        if is_main_heading:
            # Se temos heading anterior COM conteúdo, salva
            if current_heading and len('\n'.join(current_content).strip()) >= min_content_length:
                # Combina headings pendentes no título
                full_title = ' - '.join(pending_headings + [current_heading]) if pending_headings else current_heading
                doc.add_clause(
                    title=full_title,
                    content='\n'.join(current_content),
                    level=current_level,
                    source="heading"
                )
                pending_headings = []
            elif current_heading:
                # Heading sem conteúdo suficiente - acumula para combinar
                pending_headings.append(current_heading)
            
            # Novo heading principal
            current_heading = text
            current_level = heading_level
            current_content = []
        
        else:
            # TUDO que não é heading principal vira conteúdo:
            # - Sub-headings (3, 4, 5)
            # - Parágrafos normais
            # - Texto formatado
            
            if current_heading:
                # Se for um sub-heading, inclui com destaque
                if heading_level > 0:
                    current_content.append(f"\n【{text}】")  # Sub-heading marcado
                else:
                    current_content.append(text)
            elif len(text) > 30:  
                # Se não há heading ainda, cria PREÂMBULO
                if not current_heading:
                    current_heading = "PREÂMBULO"
                    current_level = 0
                    current_content = [text]
    
    # Adiciona último heading com conteúdo
    if current_heading and len('\n'.join(current_content).strip()) >= min_content_length:
        full_title = ' - '.join(pending_headings + [current_heading]) if pending_headings else current_heading
        doc.add_clause(
            title=full_title,
            content='\n'.join(current_content),
            level=current_level,
            source="heading"
        )
    
    # Processa TABELAS separadamente
    for table_idx, table in enumerate(docx_file.tables):
        rows_text = []
        for row in table.rows:
            cells_text = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if cells_text:
                rows_text.append(' | '.join(cells_text))
        
        if rows_text:
            first_row = rows_text[0][:80]
            all_content = '\n'.join(rows_text)
            doc.add_clause(
                title=f"TABELA {table_idx + 1}: {first_row}",
                content=all_content,
                level=99,  # Nível especial para tabelas
                source="table"
            )
    
    # Metadata
    doc.metadata['total_clauses'] = len(doc.clauses)
    doc.metadata['num_tables'] = len(docx_file.tables)
    
    return doc


def parse_document(filepath: str) -> Document:
    """
    Função principal de parsing - tenta V2 primeiro, fallback para V1
    """
    path = Path(filepath)
    
    if path.suffix.lower() != '.docx':
        # Se não for DOCX, usa método antigo
        from parsing import parse_document as old_parse
        return old_parse(filepath)
    
    # Tenta parsing V2 (baseado em Headings)
    doc = parse_docx_by_headings(filepath)
    
    # Se encontrou poucas cláusulas, faz fallback para V1
    if len(doc.clauses) < 5:
        print("⚠️ Poucas cláusulas com Headings, usando método tradicional...")
        from parsing import parse_document as old_parse
        return old_parse(filepath)
    
    print(f"✅ Parsing V2: {len(doc.clauses)} cláusulas baseadas em Headings")
    return doc
